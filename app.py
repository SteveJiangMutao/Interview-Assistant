import streamlit as st
import google.generativeai as genai
import tempfile
import os
import time
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
import io
import datetime
from google.api_core import retry

# --- 🔧 配置项：Logo 文件 ---
LOGO_PATH = "logo.png" 

# --- 页面配置 ---
st.set_page_config(
    page_title="Clearstate Insight Assistant",
    layout="wide",
    page_icon="🧬",
    initial_sidebar_state="expanded"
)

# --- CSS 样式 (深度间距修复) ---
st.markdown("""
<style>
    /* 全局字体优化 */
    .main-header { font-size: 2.0rem; color: #2c3e50; font-weight: bold; margin-bottom: 5px; }
    .sub-header { font-size: 1.0rem; color: #7f8c8d; margin-bottom: 20px; }
    .developer-credit { font-size: 0.85rem; color: #95a5a6; margin-top: 50px; border-top: 1px solid #bdc3c7; padding-top: 10px; }
    
    /* --- 侧边栏排版系统 --- */

    /* 1. 一级标题 (Level 1) */
    .sidebar-h1 {
        font-family: "Source Sans Pro", sans-serif;
        font-size: 18px !important;
        font-weight: 700 !important;
        color: #000000 !important;
        margin-top: 35px !important;   /* 大模块之间拉开距离 */
        margin-bottom: 10px !important; /* 标题与下方内容保持适度距离 */
        white-space: nowrap !important;
        line-height: 1.4 !important;
    }
    
    /* 2. 二级标题 (Level 2 - Label) */
    .sidebar-h2 {
        font-family: "Source Sans Pro", sans-serif;
        font-size: 15px !important;
        font-weight: 700 !important;
        color: #31333F !important;
        margin-top: 15px !important;   /* 与上一个输入框拉开距离 */
        margin-bottom: -15px !important; /* 核心：负边距，强行拉近与下方输入框的距离 */
        white-space: nowrap !important;
        line-height: 1.2 !important;
        z-index: 1; /* 确保文字在图层上方 */
        position: relative;
    }

    /* 3. 正文控件 (Radio/Input) 样式覆写 */
    
    /* Radio 按钮文字 */
    div[data-testid="stRadio"] label p {
        font-size: 14px !important;
        font-weight: 600 !important; /* 补加粗 */
        color: #31333F !important;
    }
    
    /* 输入框内部文字 */
    div[data-testid="stTextInput"] input {
        font-size: 14px !important;
        font-weight: 600 !important; /* 补加粗 */
        color: #31333F !important;
    }
    
    /* 日期选择器文字 */
    div[data-testid="stDateInput"] input {
        font-size: 14px !important;
        font-weight: 600 !important; /* 补加粗 */
    }

    /* 4. 消除 Streamlit 默认的大边距 */
    /* 这一步非常关键，去掉控件自带的 margin，完全由我们的 H1/H2 控制节奏 */
    div[data-testid="stRadio"], 
    div[data-testid="stTextInput"], 
    div[data-testid="stDateInput"] {
        margin-top: 0px !important;
        margin-bottom: 0px !important;
    }
    
    /* 针对第一个元素的特殊处理，防止顶部太挤 */
    .block-container {
        padding-top: 2rem;
    }
    
    div[data-testid="stFileUploader"] { margin-top: 20px; }
</style>
""", unsafe_allow_html=True)

# --- 辅助函数：渲染自定义侧边栏标题 ---
def render_h1(text):
    st.sidebar.markdown(f"<div class='sidebar-h1'>{text}</div>", unsafe_allow_html=True)

def render_h2(text):
    st.sidebar.markdown(f"<div class='sidebar-h2'>{text}</div>", unsafe_allow_html=True)

# --- Session State ---
if 'analysis_result' not in st.session_state:
    st.session_state['analysis_result'] = None

# --- 🧹 文本清洗函数 ---
def clean_text(text):
    """去除 Markdown 符号，保持文本纯净"""
    if isinstance(text, str):
        text = text.replace("**", "").replace("__", "")
        text = text.replace("##", "").replace("###", "")
        return text.strip()
    return text

# --- Word 格式化辅助函数 ---
def set_font_style(run, font_size=11, bold=False):
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold

def add_styled_paragraph(doc, text, bold=False, size=11, is_bullet=False, indent_level=0):
    clean_content = clean_text(str(text))
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    
    # --- 悬挂缩进逻辑 (Strict Hanging Indent) ---
    if is_bullet:
        base_indent = 0.25
        total_indent = base_indent + (indent_level * 0.25)
        
        p.paragraph_format.left_indent = Inches(total_indent)
        p.paragraph_format.first_line_indent = Inches(-base_indent)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(total_indent), WD_TAB_ALIGNMENT.LEFT)
        
        final_text = f"•\t{clean_content}"
        run = p.add_run(final_text)
    else:
        run = p.add_run(clean_content)
    
    set_font_style(run, font_size=size, bold=bold)
    return p

# --- 🌍 标题映射字典 ---
SECTION_HEADERS = {
    "commercial": {
        "zh": {
            "company_sales": "1. 厂家销售表现",
            "sales_marketing": "2. 销售与营销策略",
            "channel_strategy": "3. 销售渠道策略",
            "org_structure": "4. 组织架构与人员",
            "competition": "5. 竞争格局",
            "trends": "6. 行业趋势"
        },
        "en": {
            "company_sales": "1. Company Sales Performance",
            "sales_marketing": "2. Sales & Marketing Strategy",
            "channel_strategy": "3. Sales Channel Strategy",
            "org_structure": "4. Organizational Structure",
            "competition": "5. Competition Landscape",
            "trends": "6. Industry Trends"
        }
    },
    "clinical": {
        "zh": {
            "clinical_value": "1. 临床价值与疗效",
            "adoption": "2. 临床应用与术式",
            "competition": "3. 竞品对比",
            "pain_points": "4. 未满足需求与痛点",
            "expectations": "5. 未来预期"
        },
        "en": {
            "clinical_value": "1. Clinical Value & Efficacy",
            "adoption": "2. Adoption & Usage",
            "competition": "3. Competitive Comparison",
            "pain_points": "4. Unmet Needs & Pain Points",
            "expectations": "5. Future Expectations"
        }
    },
    "meeting": {
        "zh": {
            "meeting_context": "1. 会议背景与参会人",
            "key_discussion": "2. 核心讨论内容",
            "conclusions": "3. 结论与决策",
            "action_items": "4. 待办事项与下一步 (Follow-up)"
        },
        "en": {
            "meeting_context": "1. Context & Attendees",
            "key_discussion": "2. Key Discussion Points",
            "conclusions": "3. Conclusions & Decisions",
            "action_items": "4. Action Items & Follow-ups"
        }
    }
}

# --- Word 生成逻辑 ---
def generate_word_report(data, company, product, date, mode, meeting_topic=""):
    doc = Document()
    
    # 0. Logo (右上角, 高度 0.65cm)
    section = doc.sections[0]
    header = section.header
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if os.path.exists(LOGO_PATH):
        try:
            run_header = p_header.add_run()
            run_header.add_picture(LOGO_PATH, height=Cm(0.65))
        except Exception as e:
            print(f"Logo Error: {e}")
    
    # 语言判断
    lang = data.get('language', 'en')
    if 'zh' in lang.lower() or 'chinese' in lang.lower() or 'cn' in lang.lower():
        lang_code = 'zh'
    else:
        lang_code = 'en'

    # 1. 标题与基础信息
    if lang_code == 'zh':
        if mode == 'meeting':
            main_title = meeting_topic if meeting_topic else "内部会议"
            title_text = f"{main_title} - 会议纪要"
            type_text = '会议/讨论'
        else:
            title_text = f"{company} - {product} 访谈记录"
            type_text = '商业/厂商' if mode == 'commercial' else '临床/专家'
            
        date_prefix = "日期"
        type_prefix = "类型"
        exec_title = "摘要概览" if mode == 'meeting' else "执行摘要"
        other_title = "其他补充" if mode == 'meeting' else "其他发现"
    else:
        if mode == 'meeting':
            main_title = meeting_topic if meeting_topic else "Internal Meeting"
            title_text = f"{main_title} - Meeting Minutes"
            type_text = 'Meeting/Discussion'
        else:
            title_text = f"{company} - {product} Interview Record"
            type_text = 'Trade' if mode == 'commercial' else 'Clinical/Expert'
            
        date_prefix = "Date"
        type_prefix = "Type"
        exec_title = "Overview" if mode == 'meeting' else "Executive Summary"
        other_title = "Other Findings"

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(title_text)
    set_font_style(run_title, font_size=16, bold=True)
    
    # Meta Info
    info_text = f"{date_prefix}: {date} | {type_prefix}: {type_text}"
    add_styled_paragraph(doc, info_text, size=10.5, bold=False)
    doc.add_paragraph("-" * 80)

    # 2. Executive Summary
    summary = data.get('executive_summary', '')
    if summary:
        add_styled_paragraph(doc, exec_title, size=14, bold=True)
        add_styled_paragraph(doc, summary, size=11)

    # 3. Structured Analysis
    header_map = SECTION_HEADERS.get(mode, {}).get(lang_code, {})
    structured = data.get('structured_analysis', {})
    
    if structured:
        key_order = []
        if mode == 'commercial':
            key_order = ['company_sales', 'sales_marketing', 'channel_strategy', 'org_structure', 'competition', 'trends']
        elif mode == 'clinical':
            key_order = ['clinical_value', 'adoption', 'competition', 'pain_points', 'expectations']
        elif mode == 'meeting':
            key_order = ['meeting_context', 'key_discussion', 'conclusions', 'action_items']

        for key in key_order:
            if key in structured:
                points = structured[key]
                display_title = header_map.get(key, key.title())
                add_styled_paragraph(doc, display_title, size=12, bold=True)
                
                if isinstance(points, list):
                    for point in points:
                        add_styled_paragraph(doc, point, size=11, is_bullet=True, indent_level=0)
                else:
                    add_styled_paragraph(doc, str(points), size=11)

    # 4. Other Findings
    other_dims = data.get('other_dimensions', {})
    if other_dims:
        add_styled_paragraph(doc, other_title, size=14, bold=True)
        for k, v in other_dims.items():
            clean_k = clean_text(k)
            add_styled_paragraph(doc, clean_k, size=12, bold=True)
            if isinstance(v, list):
                for point in v:
                    add_styled_paragraph(doc, point, size=11, is_bullet=True)
            else:
                add_styled_paragraph(doc, str(v), size=11)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 核心逻辑类 ---
class InterviewAnalyzer:
    def __init__(self, api_key):
        self.api_key = api_key
        try:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel('gemini-3-pro-preview') 
        except Exception as e:
            st.error(f"API Error: {e}")

    def process_audio(self, audio_file_path):
        try:
            myfile = genai.upload_file(audio_file_path)
            with st.spinner("🎧 Uploading & Processing Audio... / 正在上传并解析音频..."):
                while myfile.state.name == "PROCESSING":
                    time.sleep(2)
                    myfile = genai.get_file(myfile.name)
            if myfile.state.name == "FAILED":
                st.error("Audio processing failed.")
                return None
            return myfile
        except Exception as e:
            st.error(f"Upload Error: {e}")
            return None

    def analyze_interview(self, audio_resource, mode):
        # 框架定义
        if mode == "commercial":
            keys_instruction = """
            Use these EXACT keys for `structured_analysis`:
            - `company_sales` (for Interviewed Manufacturer's Sales Performance)
            - `sales_marketing` (for Sales & Marketing Strategy)
            - `channel_strategy` (for Sales Channel Strategy - Distributor Focus)
            - `org_structure` (for Organizational Structure & Personnel - Internal Teams)
            - `competition` (for Competition Landscape)
            - `trends` (for Industry Trends)
            """
            framework_desc = """
            1. Company Sales Performance: Specific sales volume, revenue, and growth of the INTERVIEWED company. (Capture all numbers).
            2. Sales & Marketing Strategy: Pricing, promotion, bidding, and marketing activities.
            3. Sales Channel Strategy: **DISTRIBUTOR MANAGEMENT ONLY**. Distribution model (agency vs platform), dealer selection, dealer management policies, and channel incentives.
            4. Organizational Structure: **INTERNAL TEAMS**. Headcount, scale, and changes specifically in **Sales Dept, Marketing Dept, and Product Dept**. (e.g., "Sales team has 50 people", "Marketing expanded by 20%").
            5. Competition Landscape: Market shares of competitors, strengths/weaknesses vs competitors.
            6. Industry Trends: Policy impact, macro environment.
            """
        elif mode == "clinical":
            keys_instruction = """
            Use these EXACT keys for `structured_analysis`:
            - `clinical_value` (for Clinical Value)
            - `adoption` (for Adoption & Usage)
            - `competition` (for Competitive Comparison)
            - `pain_points` (for Unmet Needs)
            - `expectations` (for Future Expectations)
            """
            framework_desc = """
            1. Clinical Value: Efficacy, safety.
            2. Adoption & Usage: Procedure volume, indications.
            3. Competitive Comparison: Brand vs Brand.
            4. Unmet Needs: Pain points.
            5. Future Expectations: Next-gen features.
            """
        else: # meeting
            keys_instruction = """
            Use these EXACT keys for `structured_analysis`:
            - `meeting_context` (Attendees, Background)
            - `key_discussion` (Detailed discussion points, arguments made)
            - `conclusions` (What was agreed or decided)
            - `action_items` (Follow-ups, To-dos with owners)
            """
            framework_desc = """
            1. Meeting Context: List attendees and the main purpose of the meeting.
            2. Key Discussion Points: COMPREHENSIVE summary of all topics discussed. Do not miss details.
            3. Conclusions & Decisions: Clear list of decisions made.
            4. Action Items: Specific next steps, who is responsible, and deadlines if mentioned.
            """

        system_prompt = f"""
        You are a **Senior Consultant** at Clearstate.
        Task: Create a rigorous, data-driven report based on the audio.

        ### 🚨 CRITICAL INSTRUCTIONS:
        1.  **LANGUAGE CONSISTENCY**: Detect the language. 
            - If Chinese: Output ALL content in Simplified Chinese.
            - If English: Output ALL content in English.
            - **Set the `language` field in JSON to "zh" or "en".**
        2.  **NO MARKDOWN**: Do NOT use bolding marks (like **text**) in the JSON values. Output plain text only.
        
        3.  **⛔️ STRICT ENTITY HANDLING (NO TRANSLATIONS)**: 
            - **RULE**: NEVER add a translation in parentheses.
            - **WRONG**: "泰尔茂 (Terumo)", "Medtronic (美敦力)".
            - **RIGHT**: "泰尔茂", "Medtronic".
            - **EXCEPTION**: Parentheses are ONLY allowed for **Product Models** (e.g., "乐普 (NeoVas)").

        4.  **✅ PROFESSIONAL EDITING & GRAMMAR (VERY IMPORTANT)**:
            - **Fix Spoken Errors**: Audio often contains broken grammar, slips of the tongue, or awkward phrasing.
            - **CORRECTION REQUIRED**: You MUST correct these into standard, professional written language based on context.
            - **Example**: Change "年轻患者、这害怕金属植入物患者" to "年轻患者及对金属植入物有顾虑的患者".
            - **Goal**: The output must read like a polished consulting report, not a raw transcript.

        5.  **COMPREHENSIVENESS**: 
            - For Interviews: Capture every number and logic.
            - For Meetings: **Do not omit any discussion points or follow-ups.**

        ### FRAMEWORK KEYS:
        {keys_instruction}

        ### FRAMEWORK DETAILS:
        {framework_desc}

        ### OUTPUT JSON:
        {{
            "language": "zh", 
            "executive_summary": "High-level summary...",
            "structured_analysis": {{
                "key_1": ["Point 1", "Point 2"]
            }},
            "other_dimensions": {{
                "Topic": ["Detail"]
            }}
        }}
        """
        
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]
        
        try:
            response = self.model.generate_content(
                [audio_resource, system_prompt],
                safety_settings=safety_settings,
                request_options={"timeout": 600}
            )
            
            try:
                text = response.text
                if "```json" in text:
                    text = text.replace("```json", "").replace("```", "")
                return json.loads(text.strip())
            except ValueError:
                st.error("Error: Model output was not valid JSON.")
                return None

        except Exception as e:
            st.error(f"Analysis Interrupted: {e}")
            return None

# --- UI 主程序 ---
with st.sidebar:
    st.title("Clearstate AI")
    
    st.markdown("""
    <div class='developer-credit'>
    Developed by <b>Steve Jiang</b>, Clearstate
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    api_key = st.text_input("Gemini API Key", type="password")
    
    # --- 任务模式 (Level 1) ---
    render_h1("Task Mode / 任务模式")
    
    # Select Mode (Level 2)
    render_h2("Select Mode / 选择模式")
    task_mode = st.radio(
        "Select Mode", # Hidden Label
        ("interview", "meeting"),
        format_func=lambda x: "Expert Interview (专家访谈)" if x == "interview" else "Meeting Minutes (会议纪要)",
        label_visibility="collapsed"
    )
    
    # 初始化
    company_name = ""
    product_name = ""
    meeting_topic = ""
    interview_mode = "meeting" 
    
    if task_mode == "interview":
        # --- 项目信息 (Level 1) ---
        render_h1("Project Info / 项目信息")
        
        # Company (Level 2)
        render_h2("Company / 公司名称")
        company_name = st.text_input("Company", placeholder="e.g. Medtronic", label_visibility="collapsed")
        
        # Product (Level 2)
        render_h2("Product / 产品领域")
        product_name = st.text_input("Product", placeholder="e.g. Stapler", label_visibility="collapsed")
        
        # Date (Level 2)
        render_h2("Date / 访谈日期")
        interview_date = st.date_input("Date", datetime.date.today(), label_visibility="collapsed")
        
        # --- 访谈对象 (Level 1) ---
        render_h1("Interviewee Type / 访谈对象")
        
        # Select Type (Level 2)
        render_h2("Select Type / 选择类型")
        interview_sub_type = st.radio(
            "Select Type", # Hidden Label
            ("commercial", "clinical"),
            format_func=lambda x: "Trade (商业/厂商)" if x == "commercial" else "Clinical (临床/专家)",
            label_visibility="collapsed"
        )
        interview_mode = interview_sub_type
        
    else: # Meeting Mode
        # --- 会议信息 (Level 1) ---
        render_h1("Meeting Info / 会议信息")
        
        # Topic (Level 2)
        render_h2("Topic / 会议主题")
        meeting_topic = st.text_input("Topic", placeholder="e.g. Weekly Sync", label_visibility="collapsed")
        
        # Date (Level 2)
        render_h2("Date / 会议日期")
        interview_date = st.date_input("Date", datetime.date.today(), label_visibility="collapsed")
        interview_mode = "meeting"

    st.markdown("<br>", unsafe_allow_html=True) # Spacer
    if st.button("Reset / 重置"):
        st.session_state['analysis_result'] = None
        st.rerun()

st.markdown('<div class="main-header">智能市场洞察项目辅助工具</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Intelligent Market Insight Assistant</div>', unsafe_allow_html=True)

uploaded_file = st.file_uploader("Upload Audio / 上传录音 (MP3/M4A Recommended)", type=['mp3', 'wav', 'm4a'])

if uploaded_file and st.session_state['analysis_result'] is None:
    if not api_key:
        st.error("Please enter API Key in the sidebar. / 请在侧边栏输入 API Key。")
    else:
        valid_input = True
        if task_mode == "interview":
            if not company_name or not product_name:
                st.warning("Please fill in Company & Product info. / 请填写公司和产品信息。")
                valid_input = False
        
        if valid_input:
            st.audio(uploaded_file, format='audio/mp3')
            if st.button("Start Analysis (Gemini 3 Pro)", type="primary"):
                analyzer = InterviewAnalyzer(api_key)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name

                with st.status("AI is processing... / AI 正在处理...", expanded=True) as status:
                    st.write("Uploading audio to Gemini... / 正在上传音频...")
                    audio_resource = analyzer.process_audio(tmp_file_path)
                    
                    if audio_resource:
                        st.write("Analyzing (Model: gemini-3-pro-preview)... / 正在分析...")
                        result = analyzer.analyze_interview(audio_resource, interview_mode)
                        
                        if result:
                            st.session_state['analysis_result'] = result
                            status.update(label="Done! / 完成！", state="complete", expanded=False)
                            os.remove(tmp_file_path)
                            st.rerun()

if st.session_state['analysis_result']:
    res = st.session_state['analysis_result']
    
    st.success("Analysis Complete. Please download the report. / 分析完成，请下载报告。")
    
    file_date_str = interview_date.strftime("%Y%m%d")
    
    if task_mode == "interview":
        file_name = f"Interview_{company_name}_{product_name}_{file_date_str}.docx"
    else:
        topic_str = meeting_topic if meeting_topic else "Meeting"
        file_name = f"Minutes_{topic_str}_{file_date_str}.docx"
    
    docx_file = generate_word_report(res, company_name, product_name, interview_date, interview_mode, meeting_topic)
    
    st.download_button(
        label=f"Download Word Report / 下载 Word 报告",
        data=docx_file,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )

    st.markdown("---")
    st.markdown("### Preview / 预览")
    st.write(res.get('executive_summary'))
